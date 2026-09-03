"""
Convert markdown report to Word document using python-docx
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ASSET_DIR = Path(__file__).resolve().parent.parent / 'assets'

def set_run_font(run, font_name='宋体', font_size=10.5, bold=False):
    """Set run font properties"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    # Set east Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)

def add_heading_custom(doc, text, level=1):
    """Add heading with custom styling"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    font_name = '黑体' if level == 1 else '宋体'
    font_size = 16 if level == 1 else (14 if level == 2 else 12)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=True)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_custom(doc, text, indent=True, bold=False, font_size=10.5):
    """Add paragraph with custom styling"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name='宋体', font_size=font_size, bold=bold)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p

def add_code_block(doc, code_text):
    """Add code block styled paragraph"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(code_text)
    set_run_font(run, font_name='Consolas', font_size=9)
    return p

def add_image(doc, image_path, width_cm=14, caption=None):
    """Add image with caption"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_p.add_run(caption)
        set_run_font(run, font_name='宋体', font_size=9, bold=False)
        cap_p.paragraph_format.space_after = Pt(12)
    else:
        doc.add_paragraph()

def set_table_borders(table):
    """Add borders to table"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '666666')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tbl.insert(0, tblPr)

def add_table_custom(doc, headers, rows, col_widths=None):
    """Add styled table without using Table Grid style (avoids w:shd issues)"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)
    
    # Add borders manually
    set_table_borders(table)
    
    # Header row - use bold + dark blue text
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, font_name='宋体', font_size=10, bold=True)
                run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, font_name='宋体', font_size=10)
    
    doc.add_paragraph()  # spacing after table
    return table

def main():
    output_path = sys.argv[1] if len(sys.argv) > 1 else str(Path(__file__).resolve().parent / '单基站天线方向图未知_定位研究报告.docx')
    
    doc = Document()
    
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(10.5)
    
    # Set page margins
    sections = doc.sections[0]
    sections.top_margin = Cm(2.54)
    sections.bottom_margin = Cm(2.54)
    sections.left_margin = Cm(3.17)
    sections.right_margin = Cm(3.17)
    
    # ===== COVER PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('单基站天线方向图未知')
    set_run_font(run, font_name='黑体', font_size=22, bold=True)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('定位研究报告')
    set_run_font(run, font_name='黑体', font_size=22, bold=True)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('2025年')
    set_run_font(run, font_name='宋体', font_size=14)
    
    # Page break after cover
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS (Manual) =====
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_heading.add_run('目  录')
    set_run_font(run, font_name='黑体', font_size=16, bold=True)
    doc.add_paragraph()
    
    toc_items = [
        '1. 核心问题：方向图的重要性',
        '    1.1 方向图对定位的影响',
        '    1.2 忽略方向图的后果',
        '2. 联合估计方法设计',
        '    2.1 核心思想',
        '    2.2 单扇区联合估计的优化目标',
        '    2.3 优化方法',
        '3. 实验结果与分析',
        '    3.1 宏基站场景下的定位精度',
        '    3.2 关键结论',
        '4. Android 端数据采集代码',
        '    4.1 获取 LTE 信号信息',
        '    4.2 获取 GPS 位置',
        '    4.3 权限要求',
        '5. 工程建议',
        '    5.1 数据采集策略',
        '    5.2 算法部署建议',
        '    5.3 适用场景',
        '6. 方法对比总结',
        '7. 未来工作',
    ]
    
    for title_text in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(title_text)
        set_run_font(run, font_name='宋体', font_size=12)
    
    doc.add_page_break()
    
    # ===== SECTION 1 =====
    add_heading_custom(doc, '1. 核心问题：方向图的重要性', level=1)
    
    add_heading_custom(doc, '1.1 方向图对定位的影响', level=2)
    
    add_paragraph_custom(doc, 
        '天线方向图（Radiation Pattern）描述了天线在不同方向上的辐射强度分布。'
        '在基站定位中，方向图直接决定了接收信号强度（RSSI）与方位角的关系。'
        '如果忽略方向图的影响，同样的距离可能对应完全不同的 RSSI 值，导致定位结果严重偏差。')
    
    add_table_custom(doc,
        ['问题', '答案'],
        [
            ['方向图重要吗？', '极其重要，忽略它定位完全失效'],
            ['GPR 能补偿方向图吗？', '不能，方向图是方位角函数，GPR 学不了'],
            ['不知道方向图怎么办？', '联合估计——同时恢复位置和方向图参数'],
            ['需要额外数据吗？', '需要 PCI/扇区ID 来区分不同扇区的测量点'],
        ],
        [4, 8]
    )
    
    add_heading_custom(doc, '1.2 忽略方向图的后果', level=2)
    
    add_paragraph_custom(doc,
        '当完全忽略天线方向图直接进行定位时，定位误差可达 350 米以上，'
        '这比随机猜测的结果还要差。方向图是天线辐射强度随角度变化的函数，'
        '如果不知道天线指向哪个方向，同样的距离可能对应完全不同的 RSSI 值。')
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run('忽略方向图直接定位：')
    set_run_font(run, font_name='宋体', font_size=10.5, bold=False)
    run = p.add_run('350m+ 误差')
    set_run_font(run, font_name='宋体', font_size=10.5, bold=True)
    run.font.color.rgb = RGBColor(192, 0, 0)
    run = p.add_run('（比随机猜测还差）')
    set_run_font(run, font_name='宋体', font_size=10.5, bold=False)
    
    # ===== SECTION 2 =====
    add_heading_custom(doc, '2. 联合估计方法设计', level=1)
    
    add_heading_custom(doc, '2.1 核心思想', level=2)
    
    add_paragraph_custom(doc,
        '当基站天线方向图参数（方位角、波束宽度、下倾角）全部未知时，'
        '采用联合估计策略，同时优化基站位置和方向图参数。')
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run('同时优化：')
    set_run_font(run, font_name='宋体', font_size=10.5, bold=True)
    run = p.add_run('基站位置 (x, y) + 方向图参数（方位角、波束宽度、下倾角）')
    set_run_font(run, font_name='宋体', font_size=10.5, bold=False)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run('目标函数：')
    set_run_font(run, font_name='宋体', font_size=10.5, bold=True)
    run = p.add_run('最小化预测 RSSI 与实测 RSSI 的误差')
    set_run_font(run, font_name='宋体', font_size=10.5, bold=False)
    
    add_heading_custom(doc, '2.2 单扇区联合估计的优化目标', level=2)
    
    add_paragraph_custom(doc, '单扇区联合估计的损失函数实现如下：', indent=False)
    
    code = '''def loss_single_sector(params):
    bs = params[:2]           # 基站位置 (x, y)
    azimuth = params[2]       # 方位角
    beamwidth = params[3]     # 波束宽度
    tilt = params[4]          # 下倾角
    n = params[5]             # 路径损耗指数
    
    # 计算每个测量点的预测 RSSI
    # 包含：路径损耗 + 方向图增益 + 噪声
    predicted_rssi = ...
    
    # 最小化预测与实测的均方误差
    return np.mean((predicted_rssi - measured_rssi) ** 2)'''
    
    add_code_block(doc, code)
    
    add_heading_custom(doc, '2.3 优化方法', level=2)
    
    add_paragraph_custom(doc,
        '采用 L-BFGS-B 等梯度优化算法进行参数求解。对参数边界施加约束，'
        '如波束宽度限制在 30°~120° 之间，同时采用多初值策略避免陷入局部最优。')
    
    # ===== SECTION 3 =====
    add_heading_custom(doc, '3. 实验结果与分析', level=1)
    
    add_heading_custom(doc, '3.1 宏基站场景下的定位精度', level=2)
    
    add_paragraph_custom(doc,
        '在宏基站场景下，对方向图参数全部未知的场景进行了系统性测试。'
        '测试涵盖了不同测量覆盖范围、不同采样密度的情况。', indent=False)
    
    add_table_custom(doc,
        ['测量覆盖范围', '点数', '定位误差', '方位角误差', '波束宽度误差', '可辨识性'],
        [
            ['仅主瓣 ±30°', '62', '14.6 m', '2.2°', '5.1°', '良好'],
            ['主瓣 + 一侧旁瓣', '95', '11.3 m', '1.5°', '3.8°', '优秀'],
            ['全方向（360°）', '156', '8.2 m', '0.8°', '2.1°', '极好'],
            ['稀疏采样（仅4方向）', '28', '35.4 m', '8.7°', '15.2°', '差'],
        ],
        [3.5, 1.5, 2, 2, 2.5, 2]
    )
    
    # Insert chart: positioning error by coverage
    add_image(doc, ASSET_DIR / 'chart1_coverage_error.png', width_cm=14, caption='图 1  不同测量覆盖范围下的定位精度')
    
    add_heading_custom(doc, '3.2 关键结论', level=2)
    
    add_paragraph_custom(doc, '测量覆盖范围至关重要：', indent=False, bold=True)
    add_paragraph_custom(doc,
        '仅覆盖主瓣中心 ±30° 时，定位误差约 14.6 米；'
        '当覆盖全方向时，误差可降至 8.2 米；'
        '而稀疏采样（仅4个方向）会导致算法基本失效，误差达 35.4 米。')
    
    add_paragraph_custom(doc, '方向图参数可恢复性良好：', indent=False, bold=True)
    add_paragraph_custom(doc,
        '在全方向测量条件下，方位角估计精度可达 0.8°，波束宽度估计误差仅 2.1°。'
        '这说明通过联合估计，可以在不知道先验方向图信息的情况下，'
        '较为准确地恢复出基站的位置和天线参数。')
    
    add_paragraph_custom(doc, '与已知方向图的对比：', indent=False, bold=True)
    add_table_custom(doc,
        ['条件', '定位误差'],
        [
            ['方向图已知', '5-7 m'],
            ['方向图未知（联合估计）', '8-15 m'],
            ['忽略方向图', '350m+'],
        ],
        [5, 4]
    )
    
    # Insert charts: parameter accuracy and coverage trend
    add_image(doc, ASSET_DIR / 'chart2_param_accuracy.png', width_cm=14, caption='图 2  方向图参数估计精度')
    add_image(doc, ASSET_DIR / 'chart4_coverage_trend.png', width_cm=14, caption='图 3  测量覆盖范围与定位误差关系')
    
    # ===== SECTION 4 =====
    add_heading_custom(doc, '4. Android 端数据采集代码', level=1)
    
    add_heading_custom(doc, '4.1 获取 LTE 信号信息', level=2)
    
    add_paragraph_custom(doc,
        '以下代码展示了如何在 Android 端采集 LTE 基站信号信息，'
        '包括 PCI（物理小区标识）、RSRP（参考信号接收功率）等关键参数。', indent=False)
    
    code = '''// 获取 TelephonyManager
TelephonyManager telephonyManager = 
    (TelephonyManager) getSystemService(Context.TELEPHONY_SERVICE);

// 监听信号强度变化
PhoneStateListener phoneStateListener = new PhoneStateListener() {
    @Override
    public void onSignalStrengthsChanged(SignalStrength signalStrength) {
        // 获取 LTE 信号信息
        CellInfoLte lteInfo = (CellInfoLte) cellInfo;
        CellIdentityLte identity = lteInfo.getCellIdentity();
        
        int pci = identity.getPci();              // 物理小区ID（扇区标识）
        int tac = identity.getTac();              // 跟踪区码
        int earfcn = identity.getEarfcn();        // 载波频率
        
        CellSignalStrengthLte signal = lteInfo.getCellSignalStrength();
        int rsrp = signal.getRsrp();              // 参考信号接收功率(dBm)
        int rsrq = signal.getRsrq();              // 参考信号接收质量(dB)
        int rssnr = signal.getRssnr();            // 信噪比(dB)
        
        // 记录数据：{时间, GPS位置, PCI, RSRP, RSRQ}
        logMeasurement(System.currentTimeMillis(), 
                       gpsLocation, pci, rsrp, rsrq);
    }
};

telephonyManager.listen(phoneStateListener, 
    PhoneStateListener.LISTEN_SIGNAL_STRENGTHS);'''
    
    add_code_block(doc, code)
    
    add_heading_custom(doc, '4.2 获取 GPS 位置', level=2)
    
    code = '''// 获取位置服务
LocationManager locationManager = 
    (LocationManager) getSystemService(Context.LOCATION_SERVICE);

Location loc = locationManager.getLastKnownLocation(
    LocationManager.GPS_PROVIDER);

if (loc != null) {
    double latitude = loc.getLatitude();    // 纬度
    double longitude = loc.getLongitude();  // 经度
    float accuracy = loc.getAccuracy();     // 精度（米）
}'''
    
    add_code_block(doc, code)
    
    add_heading_custom(doc, '4.3 权限要求', level=2)
    
    code = '''<uses-permission android:name="android.permission.ACCESS_FINE_LOCATION" />
<uses-permission android:name="android.permission.ACCESS_COARSE_LOCATION" />
<uses-permission android:name="android.permission.READ_PHONE_STATE" />'''
    
    add_code_block(doc, code)
    
    # ===== SECTION 5 =====
    add_heading_custom(doc, '5. 工程建议', level=1)
    
    add_heading_custom(doc, '5.1 数据采集策略', level=2)
    
    items = [
        '绕基站行走：尽量覆盖基站周围 360° 方向，至少覆盖主瓣和一侧旁瓣',
        '不同距离采样：近、中、远距离都要有测量点',
        '记录 PCI：区分不同扇区，避免数据混淆',
        'GPS 精度：确保 GPS 精度在 5 米以内，否则定位误差会被 GPS 误差主导',
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_run_font(run, font_name='宋体', font_size=10.5)
    
    add_heading_custom(doc, '5.2 算法部署建议', level=2)
    
    items = [
        '分阶段估计：第一阶段粗略估计基站位置（使用全向假设），第二阶段联合优化位置和方向图参数',
        '异常值处理：剔除 NLOS（非视距）测量点，使用 RANSAC 等鲁棒估计算法',
        '在线更新：随着数据增多，逐步 refine 估计结果，使用卡尔曼滤波或滑动窗口优化',
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_run_font(run, font_name='宋体', font_size=10.5)
    
    add_heading_custom(doc, '5.3 适用场景', level=2)
    
    add_table_custom(doc,
        ['场景', '可行性', '预期精度'],
        [
            ['单扇区宏基站 + 360°测量', '高', '8-15 m'],
            ['单扇区宏基站 + 仅主瓣测量', '中', '15-25 m'],
            ['多扇区（需区分PCI）', '高', '5-10 m'],
            ['室内微基站', '中', '需额外处理多径'],
            ['仅稀疏采样（<20点）', '低', '不可靠'],
        ],
        [5, 2.5, 3]
    )
    
    # ===== SECTION 6 =====
    add_heading_custom(doc, '6. 方法对比总结', level=1)
    
    add_paragraph_custom(doc,
        '下表对比了不同条件下基站定位方法的性能差异：', indent=False)
    
    add_table_custom(doc,
        ['方法', '前提条件', '定位误差', '优点', '缺点'],
        [
            ['固定 n + 最小二乘', '方向图已知', '5-10 m', '简单快速', '需要先验知识'],
            ['全局自适应 n', '方向图已知', '8-15 m', '适应环境', '需要拟合 n'],
            ['联合估计（本文）', '方向图未知', '8-15 m', '无需先验', '计算量大'],
            ['忽略方向图', '无', '350m+', '无', '完全失效'],
        ],
        [3.5, 3, 2, 3, 3]
    )
    
    # Insert chart: method comparison
    add_image(doc, ASSET_DIR / 'chart3_method_compare.png', width_cm=14, caption='图 4  不同定位方法误差对比')
    
    # ===== SECTION 7 =====
    add_heading_custom(doc, '7. 未来工作', level=1)
    
    items = [
        '多基站联合定位：利用多个基站信号，进一步提升精度和鲁棒性',
        '深度学习辅助：使用神经网络学习方向图与环境的复杂映射',
        '实时优化：降低计算复杂度，实现手机端实时定位',
        'NLOS 识别与补偿：区分视距和非视距测量，提高复杂环境下的精度',
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_run_font(run, font_name='宋体', font_size=10.5)
    
    # Save
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

if __name__ == '__main__':
    main()
